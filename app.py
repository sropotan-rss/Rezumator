import streamlit as st
import requests
import json
import os
import tempfile
import urllib.parse
import secrets as sec
from io import BytesIO
from PyPDF2 import PdfReader
import docx
from fpdf import FPDF
import textwrap
import datetime
import time

# ---------- КОНФИГУРАЦИЯ ----------
st.set_page_config(page_title="RSS job", layout="wide")

# --- API-ключи ---
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
YANDEX_CLIENT_ID = os.getenv("YANDEX_CLIENT_ID")
YANDEX_CLIENT_SECRET = os.getenv("YANDEX_CLIENT_SECRET")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

if not GROQ_API_KEY:
    st.error("❌ Добавь GROQ_API_KEY в Secrets (обязательно)")
    st.stop()

# --- Твой Streamlit-домен (замени, если изменится) ---
APP_DOMAIN = "https://rezumator-rm6vevs9pf5zus5bdggrjk.streamlit.app"

# --- Шрифт для PDF ---
@st.cache_resource
def get_font_path():
    url = "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans.ttf"
    try:
        r = requests.get(url, timeout=30)
        r.raise_for_status()
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.ttf')
        tmp.write(r.content)
        tmp.close()
        return tmp.name
    except Exception:
        return None
FONT_PATH = get_font_path()

# --- Стили ---
st.markdown("""
<style>
    .main { background-color: #f8fafc; }
    .main-header { font-size: 2rem; font-weight: 700; color: #1E3A8A; margin-bottom: 1rem; }
    .stButton>button {
        background-color: #F97316;
        color: white;
        border-radius: 8px;
        font-weight: 600;
        border: none;
        padding: 0.5rem 1.5rem;
    }
    .stButton>button:hover { background-color: #EA580C; }
    section[data-testid="stSidebar"] {
        background-color: #ffffff;
        border-right: 1px solid #e2e8f0;
    }
    .yandex-btn {
        background-color: #FFCC00;
        color: black;
        font-weight: bold;
        border-radius: 8px;
        padding: 0.5rem 1.5rem;
        text-decoration: none;
        display: inline-block;
    }
</style>
""", unsafe_allow_html=True)

# --- Сессия ---
if "user" not in st.session_state:
    st.session_state.user = None
if "auth_provider" not in st.session_state:
    st.session_state.auth_provider = None
if "users" not in st.session_state:
    st.session_state.users = {}
if "oauth_state" not in st.session_state:
    st.session_state.oauth_state = None

# --- Яндекс OAuth (исправленный redirect_uri) ---
def yandex_login():
    if not YANDEX_CLIENT_ID:
        return
    state = sec.token_urlsafe(16)
    st.session_state.oauth_state = state
    redirect_uri = f"{APP_DOMAIN}/oauth_callback"   # <-- ЯВНЫЙ URL
    params = {
        "response_type": "code",
        "client_id": YANDEX_CLIENT_ID,
        "redirect_uri": redirect_uri,
        "state": state,
        "force_confirm": "yes"
    }
    auth_url = f"https://oauth.yandex.ru/authorize?{urllib.parse.urlencode(params)}"
    st.markdown(f'<a href="{auth_url}" class="yandex-btn">Войти через Яндекс</a>', unsafe_allow_html=True)

def handle_yandex_callback():
    query_params = st.experimental_get_query_params()
    code = query_params.get("code")
    state = query_params.get("state")
    if code and state and state[0] == st.session_state.oauth_state:
        token_url = "https://oauth.yandex.ru/token"
        data = {
            "grant_type": "authorization_code",
            "code": code[0],
            "client_id": YANDEX_CLIENT_ID,
            "client_secret": YANDEX_CLIENT_SECRET
        }
        r = requests.post(token_url, data=data)
        if r.status_code == 200:
            token_data = r.json()
            access_token = token_data.get("access_token")
            user_info_url = "https://login.yandex.ru/info?format=json"
            headers = {"Authorization": f"OAuth {access_token}"}
            user_r = requests.get(user_info_url, headers=headers)
            if user_r.status_code == 200:
                user_info = user_r.json()
                email = user_info.get("default_email") or user_info.get("emails", [None])[0]
                if email:
                    st.session_state.user = email
                    st.session_state.auth_provider = "yandex"
                    st.experimental_set_query_params()
                    st.rerun()
        st.error("Ошибка авторизации через Яндекс")
    elif code:
        st.error("Ошибка проверки state")

# --- Локальная авторизация ---
def login(email, password):
    users = st.session_state.users
    if email in users and users[email] == password:
        st.session_state.user = email
        st.session_state.auth_provider = "local"
        return True
    return False

def register(email, password):
    users = st.session_state.users
    if email in users:
        return False
    users[email] = password
    st.session_state.users = users
    st.session_state.user = email
    st.session_state.auth_provider = "local"
    return True

def logout():
    st.session_state.user = None
    st.session_state.auth_provider = None
    st.rerun()

# --- Доступные модели ИИ ---
AVAILABLE_MODELS = {}
if GROQ_API_KEY:
    AVAILABLE_MODELS["Groq (бесплатно)"] = {
        "provider": "groq",
        "model": "llama-3.1-8b-instant",
        "api_key": GROQ_API_KEY,
        "base_url": "https://api.groq.com/openai/v1/chat/completions"
    }
if OPENAI_API_KEY:
    AVAILABLE_MODELS["OpenAI (платно)"] = {
        "provider": "openai",
        "model": "gpt-4o-mini",
        "api_key": OPENAI_API_KEY,
        "base_url": "https://api.openai.com/v1/chat/completions"
    }
if GEMINI_API_KEY:
    AVAILABLE_MODELS["Gemini (бесплатно)"] = {
        "provider": "gemini",
        "model": "gemini-2.0-flash",
        "api_key": GEMINI_API_KEY,
        "base_url": "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
    }

# --- Универсальная функция запроса к ИИ ---
def ask_ai(prompt, max_tokens=2500, retries=3):
    if not AVAILABLE_MODELS:
        return "[ОШИБКА] Нет доступных моделей"
    selected = st.session_state.get("selected_model", list(AVAILABLE_MODELS.keys())[0])
    cfg = AVAILABLE_MODELS[selected]
    provider = cfg["provider"]
    api_key = cfg["api_key"]
    model = cfg["model"]
    base_url = cfg["base_url"]

    if provider in ("groq", "openai"):
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {"model": model, "messages": [{"role": "user", "content": prompt}], "temperature": 0.7, "max_tokens": max_tokens}
        last_err = ""
        for attempt in range(retries):
            try:
                r = requests.post(base_url, headers=headers, json=payload, timeout=90)
                if r.status_code == 200:
                    return r.json()["choices"][0]["message"]["content"]
                elif r.status_code == 429:
                    last_err = f"Лимит запросов (429). Попытка {attempt+1}/{retries}"
                    time.sleep(5)
                else:
                    last_err = f"Ошибка {provider} {r.status_code}: {r.text[:200]}"
                    break
            except Exception as e:
                last_err = f"Сетевая ошибка: {e}"
                time.sleep(2)
        return f"[ОШИБКА] {last_err}"

    elif provider == "gemini":
        headers = {"Content-Type": "application/json"}
        payload = {"contents": [{"parts": [{"text": prompt}]}], "generationConfig": {"maxOutputTokens": max_tokens, "temperature": 0.7}}
        last_err = ""
        for attempt in range(retries):
            try:
                r = requests.post(f"{base_url}?key={api_key}", headers=headers, json=payload, timeout=90)
                if r.status_code == 200:
                    data = r.json()
                    return data["candidates"][0]["content"]["parts"][0]["text"]
                elif r.status_code == 429:
                    last_err = f"Лимит запросов (429). Попытка {attempt+1}/{retries}"
                    time.sleep(5)
                else:
                    last_err = f"Ошибка Gemini {r.status_code}: {r.text[:200]}"
                    break
            except Exception as e:
                last_err = f"Сетевая ошибка: {e}"
                time.sleep(2)
        return f"[ОШИБКА] {last_err}"
    return "[ОШИБКА] Неизвестный провайдер"

# --- Утилиты файлов ---
def extract_text_from_pdf(file_bytes):
    pdf = PdfReader(BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in pdf.pages)

def extract_text_from_docx(file_bytes):
    doc = docx.Document(BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs)

def create_docx(text):
    doc = docx.Document()
    doc.add_heading('RSS job', 0)
    for line in text.split('\n'):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def create_pdf(text):
    if not FONT_PATH:
        return BytesIO(text.encode('utf-8'))
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font('DejaVu', '', FONT_PATH, uni=True)
    pdf.set_font('DejaVu', '', 12)
    for line in text.split('\n'):
        for wline in textwrap.wrap(line, width=80):
            pdf.cell(0, 8, wline, ln=True)
    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf

def search_hh_vacancies(text, area=113):
    r = requests.get("https://api.hh.ru/vacancies", params={"text": text, "area": area, "per_page": 10}, headers={"User-Agent": "RSSjob/1.0"})
    return r.json().get("items", []) if r.status_code == 200 else []

# --- ИИ-функции ---
def ai_roast(resume_text):
    prompt = f"Проведи аудит резюме по 5 критериям (1-10): оформление, опыт, навыки, образование, ATS. Дай вердикт и советы.\nРезюме: {resume_text[:3000]}"
    return ask_ai(prompt, max_tokens=2000)

def ai_cover_letter(resume_text, job_desc):
    prompt = f"Напиши сопроводительное письмо (до 500 символов) для резюме и вакансии.\nРезюме: {resume_text[:2000]}\nВакансия: {job_desc[:1000]}"
    return ask_ai(prompt, max_tokens=800)

def ai_checklist():
    return "Чек-лист: 1. Проверь контакты, 2. Добавь ключевые слова, 3. Укажи достижения с цифрами."

def ai_linkedin_audit():
    return "Анализ LinkedIn: добавь хэштеги, обнови раздел 'Обо мне', получи рекомендации."

def ai_achievements(resume_text):
    prompt = f"Извлеки ключевые достижения из резюме и предложи, как их лучше описать.\nРезюме: {resume_text[:2000]}"
    return ask_ai(prompt, max_tokens=1000)

def ai_radars():
    return "Радары: мониторинг вакансий по твоим ключевым словам — в разработке."

def ai_market():
    return "Рынок: средняя зарплата операционного директора — 270 000 ₽, востребованы навыки управления цепочками поставок."

# --- Интерфейс ---
def main():
    user = st.session_state.user
    with st.sidebar:
        st.markdown("## 🚀 RSS job")
        st.markdown(f"👤 {user}")
        if AVAILABLE_MODELS:
            model_names = list(AVAILABLE_MODELS.keys())
            if "selected_model" not in st.session_state:
                st.session_state.selected_model = model_names[0]
            st.selectbox("🧠 Модель ИИ", model_names, key="selected_model")
        menu = st.radio("Меню", [
            "🏠 Платформа", "🔍 Вакансии", "📨 Отклики", "📄 Резюме",
            "🔥 Аудит", "✉️ Письмо", "✅ Чек-лист", "🔗 LinkedIn аудит",
            "🏆 Достижения", "📊 Рынок", "📡 Радары", "⚙️ Настройки"
        ], label_visibility="collapsed")
        if st.button("🚪 Выйти"):
            logout()

    user_data = st.session_state.setdefault(user, {
        "rules": [], "applications": [],
        "resume_original": "", "resume_improved": ""
    })

    if menu == "🏠 Платформа":
        st.markdown('<p class="main-header">🏠 Платформа</p>', unsafe_allow_html=True)
        col1, col2, col3 = st.columns(3)
        col1.metric("Правил", len(user_data["rules"]))
        col2.metric("Откликов", len(user_data["applications"]))
        col3.metric("Резюме", "Загружено" if user_data["resume_original"] else "Нет")
        st.subheader("Последние отклики")
        for app in user_data["applications"][-5:]:
            st.write(f"📌 {app['title']} — {app['employer']} ({app['status']})")

    elif menu == "🔍 Вакансии":
        st.markdown('<p class="main-header">🔍 Поиск вакансий (hh.ru)</p>', unsafe_allow_html=True)
        query = st.text_input("Ключевые слова")
        area = st.selectbox("Регион", [("РФ",113),("Москва",1),("СПб",2)], format_func=lambda x: x[0])
        if st.button("Найти"):
            vacs = search_hh_vacancies(query, area[1])
            if vacs:
                for v in vacs[:10]:
                    with st.expander(f"{v['name']} — {v['employer']['name'] if v.get('employer') else ''}"):
                        desc = v.get("snippet",{}).get("requirement","") + " " + v.get("snippet",{}).get("responsibility","")
                        st.write(desc[:300])
                        st.markdown(f"[Открыть на hh.ru]({v.get('alternate_url')})")
            else:
                st.warning("Ничего не найдено")

    elif menu == "📨 Отклики":
        st.markdown('<p class="main-header">📨 Мои отклики</p>', unsafe_allow_html=True)
        if not user_data["applications"]:
            st.info("Нет откликов")
        else:
            for i, app in enumerate(user_data["applications"]):
                with st.expander(f"{app['title']} — {app['employer']} ({app['status']})"):
                    st.write(app["description"])
                    if app.get("letter"):
                        st.code(app["letter"])
                    st.markdown(f"[Открыть на hh.ru]({app['url']})")
                    col1, col2, col3 = st.columns(3)
                    if col1.button("Отправил", key=f"sent_{i}"):
                        app["status"] = "Отправлен"
                        st.rerun()
                    if col2.button("Повторить", key=f"rep_{i}"):
                        app["status"] = "Повтор"
                        st.rerun()
                    if col3.button("Пропустить", key=f"skip_{i}"):
                        app["status"] = "Пропущен"
                        st.rerun()

    elif menu == "📄 Резюме":
        st.markdown('<p class="main-header">📄 Моё резюме</p>', unsafe_allow_html=True)
        uploaded = st.file_uploader("Загрузить PDF/DOCX/TXT", type=["pdf","docx","txt"])
        if uploaded:
            file_bytes = uploaded.read()
            if uploaded.type == "application/pdf":
                user_data["resume_original"] = extract_text_from_pdf(file_bytes)
            elif uploaded.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                user_data["resume_original"] = extract_text_from_docx(file_bytes)
            else:
                user_data["resume_original"] = file_bytes.decode("utf-8")
            st.success("Загружено!")
        if user_data["resume_original"]:
            with st.expander("Исходный текст"):
                st.text(user_data["resume_original"][:1000])
            if st.button("✨ Улучшить резюме"):
                with st.spinner("ИИ работает..."):
                    improved = ask_ai(f"Улучши резюме: {user_data['resume_original'][:3000]}", max_tokens=1500)
                if not improved.startswith("[ОШИБКА"):
                    user_data["resume_improved"] = improved
                    st.download_button("📥 DOCX", create_docx(improved), "rss_job_resume.docx")
                    st.download_button("📥 PDF", create_pdf(improved), "rss_job_resume.pdf")
                else:
                    st.error(improved)

    elif menu == "🔥 Аудит":
        st.markdown('<p class="main-header">🔥 Аудит резюме</p>', unsafe_allow_html=True)
        if user_data["resume_original"]:
            if st.button("Запустить аудит"):
                with st.spinner("Анализируем..."):
                    result = ai_roast(user_data["resume_original"])
                st.write(result)
        else:
            st.warning("Сначала загрузите резюме")

    elif menu == "✉️ Письмо":
        st.markdown('<p class="main-header">✉️ Сопроводительное письмо</p>', unsafe_allow_html=True)
        job_desc = st.text_area("Описание вакансии")
        if st.button("Сгенерировать"):
            if user_data["resume_original"] and job_desc:
                with st.spinner("..."):
                    letter = ai_cover_letter(user_data["resume_original"], job_desc)
                st.code(letter)
            else:
                st.warning("Нужны резюме и описание вакансии")

    elif menu == "✅ Чек-лист":
        st.markdown('<p class="main-header">✅ Чек-лист перед откликом</p>', unsafe_allow_html=True)
        st.write(ai_checklist())

    elif menu == "🔗 LinkedIn аудит":
        st.markdown('<p class="main-header">🔗 Аудит LinkedIn профиля</p>', unsafe_allow_html=True)
        st.write(ai_linkedin_audit())

    elif menu == "🏆 Достижения":
        st.markdown('<p class="main-header">🏆 Мои достижения</p>', unsafe_allow_html=True)
        if user_data["resume_original"]:
            if st.button("Проанализировать достижения"):
                with st.spinner("..."):
                    result = ai_achievements(user_data["resume_original"])
                st.write(result)
        else:
            st.warning("Нужно резюме")

    elif menu == "📊 Рынок":
        st.markdown('<p class="main-header">📊 Аналитика рынка</p>', unsafe_allow_html=True)
        st.write(ai_market())

    elif menu == "📡 Радары":
        st.markdown('<p class="main-header">📡 Радары вакансий</p>', unsafe_allow_html=True)
        st.write(ai_radars())

    elif menu == "⚙️ Настройки":
        st.markdown('<p class="main-header">⚙️ Настройки аккаунта</p>', unsafe_allow_html=True)
        st.write("Здесь будут настройки профиля, уведомлений и т.д.")

# --- Экран входа ---
def auth_screen():
    st.title("🔐 RSS job")
    if "code" in st.experimental_get_query_params():
        handle_yandex_callback()
    if YANDEX_CLIENT_ID:
        yandex_login()
    st.write("---")
    choice = st.radio("", ["Вход", "Регистрация"])
    email = st.text_input("Email")
    password = st.text_input("Пароль", type="password")
    if choice == "Вход":
        if st.button("Войти локально"):
            if login(email, password):
                st.success("Добро пожаловать!")
                st.rerun()
            else:
                st.error("Неверный email или пароль")
    else:
        if st.button("Зарегистрироваться локально"):
            if register(email, password):
                st.success("Регистрация успешна!")
                st.rerun()
            else:
                st.error("Этот email уже используется")

# --- Запуск ---
if st.session_state.user:
    main()
else:
    auth_screen()
